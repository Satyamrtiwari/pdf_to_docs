from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT


def generate_mediation_form(output_path="output/generated.docx"):
    doc = Document()

    # ================= PAGE MARGINS =================
    section = doc.sections[0]
    section.top_margin = Pt(36)
    section.bottom_margin = Pt(36)
    section.left_margin = Pt(36)
    section.right_margin = Pt(36)

    # ================= HEADING HELPER =================
    def center_bold(text):
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True

    # ================= HEADINGS =================
    center_bold("FORM ‘A’")
    center_bold("MEDIATION APPLICATION FORM")
    center_bold("[REFER RULE 3(1)]")

    p = doc.add_paragraph("Mumbai District Legal Services Authority")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph("City Civil Court, Mumbai")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    # ================= TABLE =================
    table = doc.add_table(rows=0, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


    # 🔑 CRITICAL FIX: disable autofit + force column widths
    table.autofit = False
    table.columns[0].width = Inches(0.5)   # Serial No.
    table.columns[1].width = Inches(2.5)   # Label
    table.columns[2].width = Inches(3.5)   # Value

    def add_row(c0="", c1="", c2="", merge=None):
        cells = table.add_row().cells
        cells[0].text = c0
        cells[1].text = c1
        cells[2].text = c2

        for c in cells:
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        if merge == "all":
            cells[0].merge(cells[2])
        elif merge == "right":
            cells[1].merge(cells[2])

    # ================= DETAILS OF PARTIES =================
    add_row("DETAILS OF PARTIES:", merge="all")

    # -------- Applicant --------
    add_row("1", "Name of Applicant", "{{client_name}}")
    add_row("", "Address and contact details of Applicant", "", merge="right")
    add_row(
        "1",
        "Address",
        "REGISTERED ADDRESS:\n{{branch_address}}\n\n"
        "CORRESPONDENCE BRANCH ADDRESS:\n{{branch_address}}"
    )
    add_row("", "Telephone No.", "{{mobile}}")
    add_row("", "Mobile No.", "")
    add_row("", "Email ID", "info@kslegal.co.in")

    # -------- Opposite Party --------
    add_row("2", "Name, Address and Contact details of Opposite Party:", "", merge="right")
    add_row("", "Address and contact details of Defendant/s", "", merge="right")
    add_row("", "Name", "{{customer_name}}")
    add_row(
        "",
        "Address",
        "REGISTERED ADDRESS:\n"
        "{% if address1 and address1 != \"\" %}{{address1}} {% else %} ________________ {% endif %}\n\n"
        "CORRESPONDENCE ADDRESS:\n"
        "{% if address1 and address1 != \"\" %}{{address1}} {% else %} ________________ {% endif %}"
    )
    add_row("", "Telephone No.", "")
    add_row("", "Mobile No.", "")
    add_row("", "Email ID", "")

    # ================= DETAILS OF DISPUTE =================
    add_row("DETAILS OF DISPUTE:", merge="all")
    add_row(
        "",
        "THE COMM. COURTS (PRE-INSTITUTION………SETTLEMENT) RULES, 2018",
        "",
        merge="right"
    )
    add_row(
        "",
        "Nature of disputes as per section 2(1)(c) of the Commercial Courts Act, "
        "2015 (4 of 2016):",
        "",
        merge="right"
    )

    # ================= SAVE =================
    doc.save(output_path)
    return output_path
